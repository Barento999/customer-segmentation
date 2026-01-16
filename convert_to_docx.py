"""
Convert Markdown documentation to DOCX format
Requires: pip install python-docx markdown
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    # Read the markdown file
    with open('BACKEND_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Customer Segmentation Backend', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Skip the first title (already added)
        if line.startswith('# Customer Segmentation Backend'):
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Intense Quote'
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            doc.add_paragraph(text, style='List Number')
        
        # Handle tables (basic support)
        elif '|' in line and line.strip().startswith('|'):
            # Skip table separators
            if '---' in line:
                continue
            # Add table rows as paragraphs (simplified)
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            doc.add_paragraph(' | '.join(cells))
        
        # Handle regular paragraphs
        elif line.strip():
            # Skip empty lines and special markers
            if not line.startswith('```') and not line.startswith('---'):
                p = doc.add_paragraph(line)
    
    # Save the document
    doc.save('BACKEND_DOCUMENTATION.docx')
    print("✅ Successfully created BACKEND_DOCUMENTATION.docx")
    print("📄 File location: BACKEND_DOCUMENTATION.docx")
    
except ImportError:
    print("❌ Error: python-docx library not installed")
    print("📦 Install it with: pip install python-docx")
    print("\n💡 Alternative: You can also:")
    print("   1. Open BACKEND_DOCUMENTATION.md in Microsoft Word")
    print("   2. Use an online converter: https://www.markdowntodocx.com/")
    print("   3. Use Pandoc: pandoc BACKEND_DOCUMENTATION.md -o BACKEND_DOCUMENTATION.docx")
except Exception as e:
    print(f"❌ Error: {e}")
    print("\n💡 Alternative conversion methods:")
    print("   1. Open BACKEND_DOCUMENTATION.md in Microsoft Word and save as DOCX")
    print("   2. Use Pandoc: pandoc BACKEND_DOCUMENTATION.md -o BACKEND_DOCUMENTATION.docx")
    print("   3. Use online converter: https://www.markdowntodocx.com/")
